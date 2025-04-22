import os
import logging
import PyPDF2
import docx
import re

logger = logging.getLogger(__name__)

def parse_document(file_path, file_type):
    """
    Parse a document and extract its text content.
    
    Args:
        file_path (str): Path to the document file
        file_type (str): Type of the document (pdf, docx, txt, etc.)
        
    Returns:
        str: Extracted text content
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return None
    
    try:
        logger.debug(f"Parsing document: {file_path} (type: {file_type})")
        
        # Extract text based on file type
        if file_type.lower() == 'pdf':
            return extract_text_from_pdf(file_path)
        elif file_type.lower() in ['docx', 'doc']:
            return extract_text_from_docx(file_path)
        elif file_type.lower() == 'txt':
            return extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return None
            
    except Exception as e:
        logger.error(f"Error parsing document: {e}", exc_info=True)
        return None

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n\n"
        
        # Clean up the text
        text = clean_text(text)
        logger.debug(f"Extracted {len(text)} characters from PDF with {num_pages} pages")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}", exc_info=True)
        return None

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Clean up the text
        text = clean_text(text)
        logger.debug(f"Extracted {len(text)} characters from DOCX")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}", exc_info=True)
        return None

def extract_text_from_txt(file_path):
    """Extract text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            text = file.read()
        
        # Clean up the text
        text = clean_text(text)
        logger.debug(f"Extracted {len(text)} characters from TXT file")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}", exc_info=True)
        return None

def clean_text(text):
    """Clean and normalize extracted text."""
    if not text:
        return ""
    
    # Replace multiple spaces with a single space
    text = re.sub(r'\s+', ' ', text)
    
    # Replace multiple newlines with a single newline
    text = re.sub(r'\n+', '\n', text)
    
    # Remove strange characters but keep important punctuation
    text = re.sub(r'[^\w\s.,;:!?$%&()-+=\'"\/\\]', '', text)
    
    return text.strip()
